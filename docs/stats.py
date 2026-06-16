import zipfile
from docx import Document
from pathlib import Path

p = Path(__file__).parent / "POYASNLYUVALNA_ZAPYSKA_UA.docx"
with zipfile.ZipFile(p) as z:
    x = z.read("word/document.xml").decode("utf-8", errors="replace")
    print("page breaks:", x.count('w:type="page"'))
    print("exact line rule:", x.count('w:lineRule="exact"'))
d = Document(str(p))
words = sum(len(p.text.split()) for p in d.paragraphs if p.text.strip())
figs = sum(1 for p in d.paragraphs if p.text.strip() == " ")
print("words:", words)
print("figure placeholders:", figs)
# estimate: 300 words/page text + 0.5 page per figure
est = words / 280 + figs * 0.45 + 8
print("estimated pages:", round(est))
