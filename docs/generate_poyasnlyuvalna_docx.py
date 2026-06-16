"""Generate formatted .docx for Google Docs / Word (Timur Yamchuk individual report)."""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

ROOT = Path(__file__).resolve().parent
MD_FILE = ROOT / "POYASNLYUVALNA_ZAPYSKA_UA.md"
OUT_DOCS = ROOT / "POYASNLYUVALNA_ZAPYSKA_UA.docx"
OUT_DOCS_BUILD = ROOT / "POYASNLYUVALNA_ZAPYSKA_UA_build.docx"
OUT_BUILD_TMP = ROOT / "_build_otush.docx"
OUT_ROOT = ROOT.parent / "POYASNLYUVALNA_ZAPYSKA_UA.docx"
OUT_DESKTOP = Path.home() / "OneDrive" / "Desktop" / "Timur_Yamchuk.docx"
OUT_DESKTOP_ALT = Path.home() / "OneDrive" / "Desktop" / "Timur_Yamchuk_v2.docx"

FONT = "Times New Roman"
SIZE = Pt(14)
TOC_SIZE = Pt(14)

MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(1)

BODY_SPACE_AFTER = Pt(12)
FIGURE_PLACEHOLDER_LINES = 8
FIGURE_LINE_CM = Cm(2)
MIN_PAGES_TARGET = 35

FIRST_LINE_INDENT = Cm(1.25)

DIPLOMA_TITLE = (
    "«РОЗРОБКА ІНФОРМАЦІЙНО-КОМУНІКАЦІЙНОЇ СИСТЕМИ "
    "РЕКРУТИНГОВОГО ПРИЗНАЧЕННЯ»"
)
TITLE_AUTHOR = "________ Ямчук Тимур Четінович"
SUPERVISOR_DEGREE = "к. ф.-м. н., доцент"
SUPERVISOR_NAME = "________ Самойленко Д. М."
STUDY_GROUP = "КН-П-222"

TOC_ENTRIES = [
    ("Список авторів", 2),
    ("Реферат", 3),
    ("Зміст", 5),
    ("Скорочення та умовні позначення", 6),
    ("Вступ", 7),
    ("РОЗДІЛ 1. Аналіз предметної області та постановка задачі", 11),
    ("1.1. Характеристика систем професійного нетворкінгу", 11),
    ("1.2. Порівняльний аналіз існуючих рішень", 14),
    ("1.3. Обґрунтування вимог до розроблюваної системи", 18),
    ("1.4. Постановка задачі", 21),
    ("РОЗДІЛ 2. Проектування та реалізація інформаційної системи", 24),
    ("2.1. Вибір архітектурного підходу та технологічного стеку", 24),
    ("2.2. Синтез структурних та інформаційних моделей", 28),
    ("2.3. Загальна архітектура клієнтської частини", 32),
    ("2.4. Реалізація модуля автентифікації", 35),
    ("2.5. Реалізація модуля вакансій", 39),
    ("2.6. Реалізація модуля повідомлень (чати)", 43),
    ("2.7. Реалізація особистого кабінету та дизайну інтерфейсу", 47),
    ("2.8. Засоби автентифікації та захисту інформації", 51),
    ("РОЗДІЛ 3. Розгортання та випробування системи", 54),
    ("3.1. Розгортання та інтеграція frontend з backend", 54),
    ("3.2. Результати інтеграційного та ручного тестування", 58),
    ("3.3. Обмеження та перспективи розвитку", 62),
    ("Висновки та рекомендації", 65),
    ("Перелік джерел", 68),
    ("Додатки", 72),
]


def set_run_font(run, bold=False, size=None):
    run.font.name = FONT
    run.font.size = size or SIZE
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def style_paragraph(
    p,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_cm=None,
    space_after=BODY_SPACE_AFTER,
    space_before=0,
    exact_line=True,
):
    pf = p.paragraph_format
    pf.alignment = align
    if exact_line:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.0
    pf.space_after = space_after
    pf.space_before = space_before
    if first_line_cm == 0:
        pf.first_line_indent = Cm(0)
    elif first_line_cm is None:
        pf.first_line_indent = FIRST_LINE_INDENT
    else:
        pf.first_line_indent = Cm(first_line_cm)


def add_para(doc, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, center=False, first_line_cm=None, space_after=BODY_SPACE_AFTER, space_before=0, in_cell=None):
    if center:
        align = WD_ALIGN_PARAGRAPH.CENTER
        first_line_cm = 0
    container = in_cell if in_cell is not None else doc
    p = container.add_paragraph() if in_cell is not None else doc.add_paragraph()
    style_paragraph(p, align=align, first_line_cm=first_line_cm, space_after=space_after, space_before=space_before)
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    return p


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tblPr.append(borders)


def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        tcBorders.append(elem)
    tcPr.append(tcBorders)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    # Remove default docGrid — it compresses line spacing in Word
    sectPr = section._sectPr
    for dg in sectPr.findall(qn("w:docGrid")):
        sectPr.remove(dg)
    return doc


def build_title_page(doc):
    """Single page: header top, Odessa bottom."""
    section = doc.sections[0]
    usable = section.page_height - section.top_margin - section.bottom_margin

    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    remove_table_borders(table)

    top_cell = table.rows[0].cells[0]
    bottom_cell = table.rows[1].cells[0]
    remove_cell_borders(top_cell)
    remove_cell_borders(bottom_cell)

    top_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    bottom_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    # ~82% top block, ~18% bottom line
    table.rows[0].height = Twips(int(usable * 0.82))
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[1].height = Twips(int(usable * 0.18))
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    ls = WD_LINE_SPACING.SINGLE

    for t in [
        "Приватний заклад вищої освіти",
        "Одеський технологічний університет «ШАГ»",
        "Кафедра інформаційних технологій та фундаментальної підготовки",
    ]:
        add_para(top_cell, t, center=True, in_cell=top_cell, space_after=0)

    add_para(top_cell, "", center=True, in_cell=top_cell, space_after=6)
    add_para(top_cell, "ПОЯСНЮВАЛЬНА ЗАПИСКА", bold=True, center=True, in_cell=top_cell, space_after=3)
    add_para(top_cell, "до дипломного проєкту (роботи)", center=True, in_cell=top_cell, space_after=6)
    add_para(top_cell, DIPLOMA_TITLE, bold=True, center=True, in_cell=top_cell, space_after=6)

    for t in [
        "на здобуття бакалаврського ступеня",
        "першого рівня вищої освіти",
        "зі спеціальності «F3 Комп'ютерні науки»",
    ]:
        add_para(top_cell, t, center=True, in_cell=top_cell, space_after=0)

    add_para(top_cell, "", center=True, in_cell=top_cell, space_after=12)
    add_para(top_cell, "Виконав:", center=True, in_cell=top_cell, space_after=0)
    add_para(top_cell, f"студент групи {STUDY_GROUP}", center=True, in_cell=top_cell, space_after=6)
    add_para(top_cell, TITLE_AUTHOR, center=True, in_cell=top_cell, space_after=18)

    add_para(top_cell, "Керівник:", center=True, in_cell=top_cell, space_after=0)
    add_para(top_cell, SUPERVISOR_DEGREE, center=True, in_cell=top_cell, space_after=6)
    add_para(top_cell, SUPERVISOR_NAME, center=True, in_cell=top_cell, space_after=0)

    add_para(bottom_cell, "Одеса – 2026", center=True, in_cell=bottom_cell, space_before=0)

    doc.add_page_break()


def add_heading(doc, text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line_cm=0, space_after=6, space_before=6)
    set_run_font(p.add_run(text.upper()), bold=True)


def add_subheading(doc, text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0, space_after=6, space_before=6)
    set_run_font(p.add_run(text), bold=True)


def add_figure_placeholder(doc):
    """Reserve space under figure caption (~1/3 page) for screenshots."""
    for _ in range(FIGURE_PLACEHOLDER_LINES):
        p = doc.add_paragraph()
        style_paragraph(
            p,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_cm=0,
            space_before=FIGURE_LINE_CM,
            space_after=FIGURE_LINE_CM,
            exact_line=False,
        )
        p.add_run(" ")


def subsection_page_break(title):
    if not title[0].isdigit():
        return False
    parts = title.split(".", 1)
    if len(parts) != 2 or not parts[0].isdigit():
        return False
    major = int(parts[0])
    if major not in (1, 2, 3):
        return False
    if title.startswith(("1.1", "1.2", "1.3", "1.4")):
        return True
    if major == 2 and title.startswith("2."):
        return True
    if major == 3 and title.startswith("3."):
        return True
    return False


def add_caption(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0, space_after=3, space_before=6)
    set_run_font(p.add_run(text), bold=False)


def add_toc_table(doc):
    add_heading(doc, "ЗМІСТ", page_break=True)

    table = doc.add_table(rows=len(TOC_ENTRIES), cols=2)
    table.autofit = False
    remove_table_borders(table)

    content_width = doc.sections[-1].page_width - doc.sections[-1].left_margin - doc.sections[-1].right_margin
    table.columns[0].width = Twips(int(content_width * 0.92))
    table.columns[1].width = Twips(int(content_width * 0.08))

    for i, (title, page) in enumerate(TOC_ENTRIES):
        left = table.rows[i].cells[0]
        right = table.rows[i].cells[1]
        remove_cell_borders(left)
        remove_cell_borders(right)

        lp = left.paragraphs[0]
        style_paragraph(lp, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0, space_after=0, exact_line=False)
        set_run_font(lp.add_run(title), size=TOC_SIZE)

        rp = right.paragraphs[0]
        style_paragraph(rp, align=WD_ALIGN_PARAGRAPH.RIGHT, first_line_cm=0, space_after=0, exact_line=False)
        set_run_font(rp.add_run(str(page)), size=TOC_SIZE)


def parse_table_rows(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if set(line.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
    return rows


def add_table(doc, rows, caption=None):
    if caption:
        add_caption(doc, caption)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        tr = table.rows[i]
        tr.height = Cm(2)
        tr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            for p in cell.paragraphs:
                style_paragraph(
                    p,
                    align=WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_cm=0,
                    space_after=Pt(6),
                    exact_line=True,
                )
                for run in p.runs:
                    set_run_font(run, bold=(i == 0), size=Pt(12))
    add_para(doc, "", space_after=Pt(12), first_line_cm=0)


def is_major_heading(line):
    u = line.strip().upper()
    return u in {
        "СПИСОК АВТОРІВ",
        "РЕФЕРАТ",
        "ЗМІСТ",
        "СКОРОЧЕННЯ ТА УМОВНІ ПОЗНАЧЕННЯ",
        "ВСТУП",
        "ВИСНОВКИ ТА РЕКОМЕНДАЦІЇ",
        "ПЕРЕЛІК ДЖЕРЕЛ",
        "ДОДАТКИ",
    } or u.startswith("РОЗДІЛ ")


def is_subheading(line):
    s = line.strip()
    if len(s) < 4 or s.startswith("|"):
        return False
    parts = s.split(".", 1)
    if len(parts) != 2:
        return False
    num, rest = parts
    nums = num.split(".")
    return all(n.isdigit() for n in nums if n) and bool(rest.strip())


def skip_toc_block(lines, start_index):
    i = start_index
    while i < len(lines):
        s = lines[i].strip()
        if not s or s == "---":
            i += 1
            continue
        if s.upper() == "ЗМІСТ" or s.startswith("<!--"):
            i += 1
            continue
        if is_major_heading(s) and s.upper() != "ЗМІСТ":
            break
        i += 1
    return i


def build_body(doc, text):
    lines = text.splitlines()
    i = 0
    pending_caption = None

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---" or stripped.startswith("<!--"):
            i += 1
            continue
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table(doc, parse_table_rows(block), caption=pending_caption)
            pending_caption = None
            continue

        if stripped.upper() == "ЗМІСТ":
            add_toc_table(doc)
            i = skip_toc_block(lines, i + 1)
            continue

        if stripped.startswith("Таблиця "):
            pending_caption = stripped
            i += 1
            continue

        if stripped.startswith("Рисунок "):
            add_caption(doc, stripped)
            add_figure_placeholder(doc)
            i += 1
            continue

        if is_major_heading(stripped):
            add_heading(doc, stripped, page_break=(stripped.upper() not in {"СПИСОК АВТОРІВ", "РЕФЕРАТ"}))
            i += 1
            continue

        if is_subheading(stripped):
            add_subheading(doc, stripped, page_break=subsection_page_break(stripped))
            i += 1
            continue

        if stripped.startswith("Лістинг "):
            add_caption(doc, stripped)
            i += 1
            continue

        if stripped.startswith("export ") or stripped.startswith("function ") or (stripped and stripped[0].isdigit() and ". " in stripped[:4]):
            # code / numbered algorithm lines
            p = doc.add_paragraph()
            style_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_cm=0, space_after=0, exact_line=False)
            run = p.add_run(stripped)
            set_run_font(run, size=Pt(12))
            i += 1
            continue

        if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == ".":
            add_para(doc, stripped, first_line_cm=0)
            i += 1
            continue

        if stripped[0:2] in {"1.", "2.", "3.", "4."} and "—" in stripped:
            add_para(doc, stripped, first_line_cm=0)
            i += 1
            continue

        add_para(doc, stripped)
        i += 1


def main():
    if not MD_FILE.exists():
        raise FileNotFoundError(MD_FILE)

    text = MD_FILE.read_text(encoding="utf-8")
    doc = setup_document()
    build_title_page(doc)

    idx = text.upper().find("СПИСОК АВТОРІВ")
    build_body(doc, text[idx:] if idx >= 0 else text)

    doc.save(OUT_BUILD_TMP)
    for src, targets in ((OUT_BUILD_TMP, (OUT_DOCS_BUILD, OUT_DOCS, OUT_ROOT, OUT_DESKTOP)),):
        for target in targets:
            try:
                import shutil
                shutil.copy2(src, target)
                print(f"Saved: {target}")
            except PermissionError:
                alt = target.with_name(target.stem + "_35p" + target.suffix)
                shutil.copy2(src, alt)
                print(f"Saved (locked, alt): {alt}")

    # Optional Word post-format (slow on large docs; generator uses 1.5 spacing)
    # subprocess.run(["python", str(ROOT / "fix_spacing.py")], check=False, cwd=str(ROOT))


if __name__ == "__main__":
    main()
